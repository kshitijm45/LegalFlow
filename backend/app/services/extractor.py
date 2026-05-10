"""
Text extraction from PDF, DOCX, and plain-text files.
Uses PyMuPDF (fitz) for PDFs and python-docx for DOCX.
"""
from __future__ import annotations

import asyncio
import io
from typing import NamedTuple


class ExtractionResult(NamedTuple):
    text: str
    page_count: int


async def extract_text(file_bytes: bytes, file_type: str) -> ExtractionResult:
    """Extract plain text from a document. Runs in a thread pool."""
    return await asyncio.to_thread(_extract_sync, file_bytes, file_type)


def _extract_sync(file_bytes: bytes, file_type: str) -> ExtractionResult:
    ft = file_type.lower().strip(".")
    if ft == "pdf":
        return _from_pdf(file_bytes)
    if ft in ("docx", "doc"):
        return _from_docx(file_bytes)
    if ft == "txt":
        return _from_txt(file_bytes)
    raise ValueError(f"Unsupported file type: {file_type}")


def _from_pdf(data: bytes) -> ExtractionResult:
    import fitz  # PyMuPDF

    doc = fitz.open(stream=data, filetype="pdf")
    pages = []
    for page in doc:
        pages.append(page.get_text())
    doc.close()
    return ExtractionResult(text="\n\n".join(pages), page_count=len(pages))


def _from_docx(data: bytes) -> ExtractionResult:
    from docx import Document

    doc = Document(io.BytesIO(data))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also extract table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return ExtractionResult(text="\n\n".join(paragraphs), page_count=1)


def _from_txt(data: bytes) -> ExtractionResult:
    text = data.decode("utf-8", errors="replace")
    return ExtractionResult(text=text, page_count=1)
